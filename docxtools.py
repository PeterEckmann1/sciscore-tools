import docx
import statistics
import subprocess
import os
import re


class Document:
    def __init__(self, file):
        if file.endswith('.doc'):
            subprocess.call(['soffice', '--headless', '--convert-to', 'docx', file])
            docx_file = file.split('/')[-1] + 'x'
            self.doc = docx.Document(docx_file)
            os.remove(docx_file)
        else:
            self.doc = docx.Document(file)
        hashes = []
        for para in self.doc.paragraphs:
            for run in para.runs:
                hashes.append(self._get_hash(run))
        self.body_hash = statistics.mode(hashes)
        self.header_hash = None

    def _get_hash(self, run):
        return hash((run.font.size, run.font.name, run.bold, run.italic, run.font.color.rgb))

    def get_text(self, section):
        text = ''
        for para in self.doc.paragraphs:
            for run in para.runs:
                text += run.text
        return text.replace('\n', ' ')


f = open('full_texts.txt', 'w', encoding='utf-8')
for f_name in os.listdir('Umairs sample files'):
    if '.doc' in f_name or '.docx' in f_name:
        f.write(f_name + ','.join(re.findall('(METHOD|Method)s?\W+[A-Z]', Document(f'Umairs sample files/{f_name}').get_text('methods'))) + '\n')